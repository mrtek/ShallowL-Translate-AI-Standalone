# ============================================================================
# ShallowL Translate AI Standalone
# AI-powered translation tool with local LLM support
# Инструмент для перевода с поддержкой локальных LLM моделей
# ============================================================================

import customtkinter as ctk
import tkinter as tk
from tkinter import filedialog, messagebox
import threading
import os
import ctypes
import sys
import argparse
import re
import json
import queue
import subprocess
import time
import requests
import docx
import fitz  # PyMuPDF
from openai import OpenAI
from collections import OrderedDict
from concurrent.futures import ThreadPoolExecutor, as_completed
import logging

# Logging setup / Настройка логирования в файл
BASE_PATH = os.path.dirname(os.path.abspath(__file__))
logging.basicConfig(
    level=logging.DEBUG,
    format='%(asctime)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler(os.path.join(BASE_PATH, 'app.log'), encoding='utf-8', mode='w'),
        logging.StreamHandler()
    ]
)

# ============================================================================
# Constants / Константы
# ============================================================================
BASE_PATH = os.path.dirname(os.path.abspath(__file__))
PROMPTS_FILE = os.path.join(BASE_PATH, "prompts.json")
CONFIG_FILE = os.path.join(BASE_PATH, "config.json")
SUPPORTED_LANGS = ["English", "Russian", "German", "French", "Spanish", "Italian", "Chinese", "Japanese"]

# Precompiled regex for text validation (performance optimization)
# Предкомпилированный regex для проверки текста (оптимизация производительности)
TEXT_CHECK_REGEX = re.compile(r'[a-zA-Zа-яА-ЯёЁ]')

# Load UI translations from JSON file / Загрузка локализации из файла
TRANSLATIONS_FILE = os.path.join(BASE_PATH, "lang", "translations.json")
try:
    with open(TRANSLATIONS_FILE, 'r', encoding='utf-8') as f:
        TRANSLATIONS = json.load(f)
except:
    # Fallback if file not found / Резервный вариант если файл не найден
    TRANSLATIONS = {"EN": {}, "RU": {}}

# ============================================================================
# Available AI Models / Доступные AI модели
# Models are organized by VRAM requirements
# Модели организованы по требованиям к видеопамяти
# ============================================================================
AVAILABLE_MODELS = {
    "━━━ Для VRAM 8GB+ ━━━": {
        "file": "",
        "url": "",
        "size": "───────",
        "speed": "",
        "quality": ""
    },
    "Dolphin Mistral Nemo 12B Q4 (Uncensored)": {
        "file": "dolphin-2.9.3-mistral-nemo-12b.Q4_K_M.gguf",
        "url": "https://huggingface.co/dphn/dolphin-2.9.3-mistral-nemo-12b-gguf/resolve/main/dolphin-2.9.3-mistral-nemo-12b.Q4_K_M.gguf",
        "size": "7.2 GB",
        "speed": "⭐⭐⭐",
        "quality": "⭐⭐⭐⭐⭐"
    },
    "Llama 3.1 8B Q6_K (Сбалансированная)": {
        "file": "Meta-Llama-3.1-8B-Instruct-Q6_K.gguf",
        "url": "https://huggingface.co/bartowski/Meta-Llama-3.1-8B-Instruct-GGUF/resolve/main/Meta-Llama-3.1-8B-Instruct-Q6_K.gguf",
        "size": "6.6 GB",
        "speed": "⭐⭐⭐⭐",
        "quality": "⭐⭐⭐⭐"
    },
    "━━━ Для VRAM 6GB+ ━━━": {
        "file": "",
        "url": "",
        "size": "───────",
        "speed": "",
        "quality": ""
    },
    "Qwen 2.5 7B Q6_K (Быстрая, мультиязычная)": {
        "file": "Qwen2.5-7B-Instruct-Q6_K_L.gguf",
        "url": "https://huggingface.co/bartowski/Qwen2.5-7B-Instruct-GGUF/resolve/main/Qwen2.5-7B-Instruct-Q6_K_L.gguf",
        "size": "6.2 GB",
        "speed": "⭐⭐⭐⭐⭐",
        "quality": "⭐⭐⭐⭐"
    },
    "Mistral 7B v0.3 Q6_K (Классика)": {
        "file": "Mistral-7B-Instruct-v0.3-Q6_K.gguf",
        "url": "https://huggingface.co/bartowski/Mistral-7B-Instruct-v0.3-GGUF/resolve/main/Mistral-7B-Instruct-v0.3-Q6_K.gguf",
        "size": "5.9 GB",
        "speed": "⭐⭐⭐⭐",
        "quality": "⭐⭐⭐⭐"
    },
    "Qwen 2.5 7B Q5_K_M (Оптимальная)": {
        "file": "Qwen2.5-7B-Instruct-Q5_K_M.gguf",
        "url": "https://huggingface.co/bartowski/Qwen2.5-7B-Instruct-GGUF/resolve/main/Qwen2.5-7B-Instruct-Q5_K_M.gguf",
        "size": "5.4 GB",
        "speed": "⭐⭐⭐⭐⭐",
        "quality": "⭐⭐⭐⭐"
    },
    "━━━ Для VRAM 4GB+ ━━━": {
        "file": "",
        "url": "",
        "size": "───────",
        "speed": "",
        "quality": ""
    },
    "Mistral 7B v0.3 Q3_K_M (Компактная)": {
        "file": "Mistral-7B-Instruct-v0.3-Q3_K_M.gguf",
        "url": "https://huggingface.co/bartowski/Mistral-7B-Instruct-v0.3-GGUF/resolve/main/Mistral-7B-Instruct-v0.3-Q3_K_M.gguf",
        "size": "3.3 GB",
        "speed": "⭐⭐⭐⭐⭐",
        "quality": "⭐⭐⭐"
    },
    "Qwen 2.5 7B Q2_K (Легкая)": {
        "file": "Qwen2.5-7B-Instruct-Q2_K.gguf",
        "url": "https://huggingface.co/bartowski/Qwen2.5-7B-Instruct-GGUF/resolve/main/Qwen2.5-7B-Instruct-Q2_K.gguf",
        "size": "2.8 GB",
        "speed": "⭐⭐⭐⭐⭐",
        "quality": "⭐⭐"
    },
    "Llama 3.2 3B Q6_K (Маленькая, быстрая)": {
        "file": "Llama-3.2-3B-Instruct-Q6_K.gguf",
        "url": "https://huggingface.co/bartowski/Llama-3.2-3B-Instruct-GGUF/resolve/main/Llama-3.2-3B-Instruct-Q6_K.gguf",
        "size": "2.4 GB",
        "speed": "⭐⭐⭐⭐⭐",
        "quality": "⭐⭐⭐"
    },
    "Gemma 2 2B Q8_0 (Компактная от Google)": {
        "file": "gemma-2-2b-it-Q8_0.gguf",
        "url": "https://huggingface.co/bartowski/gemma-2-2b-it-GGUF/resolve/main/gemma-2-2b-it-Q8_0.gguf",
        "size": "2.7 GB",
        "speed": "⭐⭐⭐⭐⭐",
        "quality": "⭐⭐⭐"
    }
}

# ============================================================================
# Optimal batch presets for each model / Оптимальные пресеты батчинга для каждой модели
# Higher tokens/lines = faster but may reduce quality
# Больше токенов/строк = быстрее, но может снизить качество
# ============================================================================
MODEL_BATCH_PRESETS = {
    "dolphin-2.9.3-mistral-nemo-12b-Q4_K_M.gguf": {"tokens": 500, "lines": 6, "vram": "8GB+"},
    "Meta-Llama-3.1-8B-Instruct-Q6_K.gguf": {"tokens": 600, "lines": 8, "vram": "8GB+"},
    "Qwen2.5-7B-Instruct-Q6_K_L.gguf": {"tokens": 700, "lines": 10, "vram": "6GB+"},
    "Mistral-7B-Instruct-v0.3-Q6_K.gguf": {"tokens": 650, "lines": 9, "vram": "6GB+"},
    "Qwen2.5-7B-Instruct-Q5_K_M.gguf": {"tokens": 750, "lines": 10, "vram": "6GB+"},
    "Mistral-7B-Instruct-v0.3-Q3_K_M.gguf": {"tokens": 900, "lines": 12, "vram": "4GB+"},
    "Qwen2.5-7B-Instruct-Q2_K.gguf": {"tokens": 1000, "lines": 15, "vram": "4GB+"},
    "Llama-3.2-3B-Instruct-Q6_K.gguf": {"tokens": 1100, "lines": 15, "vram": "4GB+"},
    "gemma-2-2b-it-Q8_0.gguf": {"tokens": 1200, "lines": 18, "vram": "4GB+"}
}

# ============================================================================
# Default translation prompts / Дефолтные промпты для перевода
# Each prompt has: text, temperature, top_p, penalty parameters
# Каждый промпт содержит: текст, температуру, top_p, штраф за повторения
# ============================================================================
DEFAULT_PROMPTS = {
    "Строгий / Strict": {"text": "You are a professional native {target} translator. Translate from {source} to {target} with maximum accuracy. Output ONLY the translation.", "temp": 0.1, "top_p": 0.9, "penalty": 1.1},
    "Литературный / Literary": {"text": "You are a creative writer. Translate from {source} to {target} beautifully, adapting metaphors. Output ONLY the translation.", "temp": 0.7, "top_p": 0.95, "penalty": 1.15},
    "Технический / Technical": {"text": "You are a technical expert. Translate using precise professional terminology. Output ONLY the translation.", "temp": 0.1, "top_p": 0.5, "penalty": 1.1},
    "От женского лица / Female Speaker": {"text": "You are translating from {source} to {target}. The speaker is FEMALE. Use feminine gender endings (e.g. 'Я пошла'). Output ONLY the translation.", "temp": 0.3, "top_p": 0.9, "penalty": 1.1},
    "От мужского лица / Male Speaker": {"text": "You are translating from {source} to {target}. The speaker is MALE. Use masculine gender endings (e.g. 'Я пошел'). Output ONLY the translation.", "temp": 0.3, "top_p": 0.9, "penalty": 1.1},
    "Перевод игр / Game Localization": {"text": "You are a game localizer translating from {source} to {target}. Preserve all tags, placeholders (like {0}, %s), and special symbols. Keep the tone engaging. Output ONLY the translation.", "temp": 0.4, "top_p": 0.9, "penalty": 1.15}
}

# Default application config / Конфигурация приложения по умолчанию
DEFAULT_CONFIG = {"theme": "dark", "lang": "RU", "batch_tokens": 600, "batch_lines": 8, "smart_mode": True}

# ============================================================================
# Prompt Manager Window / Окно менеджера промптов
# Allows editing translation prompts and AI parameters
# Позволяет редактировать промпты перевода и параметры AI
# ============================================================================
class PromptManager(ctk.CTkToplevel):
    def __init__(self, parent, prompts, save_callback):
        super().__init__(parent)
        self.title("⚙ Prompt Editor / Редактор промптов")
        self.geometry("900x700")
        self.attributes("-topmost", True)
        self.prompts, self.save_callback = prompts, save_callback

        # Главный контейнер с двумя колонками
        main_container = ctk.CTkFrame(self, fg_color="transparent")
        main_container.pack(fill="both", expand=True, padx=10, pady=10)
        main_container.grid_columnconfigure(0, weight=1)
        main_container.grid_columnconfigure(1, weight=1)
        main_container.grid_rowconfigure(0, weight=1)

        # Левая колонка - список и редактор
        left_col = ctk.CTkFrame(main_container)
        left_col.grid(row=0, column=0, sticky="nsew", padx=(0, 5))
        left_col.grid_columnconfigure(0, weight=1)
        left_col.grid_rowconfigure(1, weight=1)

        # Список промптов
        list_label = ctk.CTkLabel(left_col, text="Prompts List:", font=("Arial", 12, "bold"))
        list_label.grid(row=0, column=0, sticky="w", padx=10, pady=(10, 5))

        self.listbox = tk.Listbox(left_col, bg="#2b2b2b", fg="white", font=("Arial", 10),
                                  borderwidth=0, selectbackground="#1f538d", height=8)
        self.listbox.grid(row=1, column=0, sticky="nsew", padx=10, pady=5)
        self.listbox.bind("<<ListboxSelect>>", self.on_select)

        # Редактор промпта
        edit_label = ctk.CTkLabel(left_col, text="Edit Prompt:", font=("Arial", 12, "bold"))
        edit_label.grid(row=2, column=0, sticky="w", padx=10, pady=(10, 5))

        self.name_entry = ctk.CTkEntry(left_col, placeholder_text="Prompt Name", font=("Arial", 12))
        self.name_entry.grid(row=3, column=0, sticky="ew", padx=10, pady=5)

        self.text_editor = ctk.CTkTextbox(left_col, height=150, font=("Arial", 10))
        self.text_editor.grid(row=4, column=0, sticky="nsew", padx=10, pady=5)
        left_col.grid_rowconfigure(4, weight=1)

        # Слайдеры параметров
        sliders_frame = ctk.CTkFrame(left_col)
        sliders_frame.grid(row=5, column=0, sticky="ew", padx=10, pady=5)

        ctk.CTkLabel(sliders_frame, text="Temperature (0.1-1.2):", font=("Arial", 10)).pack(anchor="w", padx=5)
        temp_row = ctk.CTkFrame(sliders_frame, fg_color="transparent")
        temp_row.pack(fill="x", padx=5)
        self.s_temp = ctk.CTkSlider(temp_row, from_=0.1, to=1.2,
                                    command=lambda v: self.l_temp.configure(text=f"{v:.2f}"))
        self.s_temp.pack(side="left", fill="x", expand=True)
        self.l_temp = ctk.CTkLabel(temp_row, text="0.2", width=40)
        self.l_temp.pack(side="left", padx=5)

        ctk.CTkLabel(sliders_frame, text="Top-P (0.1-1.0):", font=("Arial", 10)).pack(anchor="w", padx=5, pady=(5,0))
        top_row = ctk.CTkFrame(sliders_frame, fg_color="transparent")
        top_row.pack(fill="x", padx=5)
        self.s_top = ctk.CTkSlider(top_row, from_=0.1, to=1.0,
                                   command=lambda v: self.l_top.configure(text=f"{v:.2f}"))
        self.s_top.pack(side="left", fill="x", expand=True)
        self.l_top = ctk.CTkLabel(top_row, text="0.9", width=40)
        self.l_top.pack(side="left", padx=5)

        ctk.CTkLabel(sliders_frame, text="Penalty (1.0-1.5):", font=("Arial", 10)).pack(anchor="w", padx=5, pady=(5,0))
        pen_row = ctk.CTkFrame(sliders_frame, fg_color="transparent")
        pen_row.pack(fill="x", padx=5)
        self.s_pen = ctk.CTkSlider(pen_row, from_=1.0, to=1.5,
                                   command=lambda v: self.l_pen.configure(text=f"{v:.2f}"))
        self.s_pen.pack(side="left", fill="x", expand=True)
        self.l_pen = ctk.CTkLabel(pen_row, text="1.1", width=40)
        self.l_pen.pack(side="left", padx=5)

        # Кнопки управления
        btn_row = ctk.CTkFrame(left_col, fg_color="transparent")
        btn_row.grid(row=6, column=0, sticky="ew", padx=10, pady=10)
        ctk.CTkButton(btn_row, text=f"➕ {parent.t('new', 'prompt_manager')}", width=80, command=self.add_p).pack(side="left", padx=3)
        ctk.CTkButton(btn_row, text=f"❌ {parent.t('delete', 'prompt_manager')}", width=80, fg_color="#dc3545", command=self.del_p).pack(side="left", padx=3)
        ctk.CTkButton(btn_row, text=f"💾 {parent.t('save', 'prompt_manager')}", width=100, fg_color="#28a745", command=self.save_p).pack(side="right", padx=3)

        # Правая колонка - справка по параметрам (двуязычная)
        right_col = ctk.CTkFrame(main_container)
        right_col.grid(row=0, column=1, sticky="nsew", padx=(5, 0))

        help_label = ctk.CTkLabel(right_col, text="📖 AI Parameters Guide" if parent.ui_lang == "EN" else "📖 Справка по параметрам ИИ",
                                  font=("Arial", 13, "bold"))
        help_label.pack(pady=10)

        help_text = ctk.CTkTextbox(right_col, font=("Arial", 10), wrap="word")
        help_text.pack(fill="both", expand=True, padx=10, pady=(0, 10))

        if parent.ui_lang == "EN":
            help_content = """🌡 TEMPERATURE
Controls creativity and randomness.

• 0.1-0.3: Very precise, deterministic
  Perfect for technical translation

• 0.4-0.6: Balanced, natural
  Good for general text

• 0.7-1.2: Creative, varied
  Best for literary translation

━━━━━━━━━━━━━━━━━━━━━━

🎯 TOP-P (Nucleus Sampling)
Limits vocabulary diversity.

• 0.1-0.5: Conservative words
  Strict, formal style

• 0.6-0.8: Balanced vocabulary
  Natural, everyday language

• 0.9-1.0: Full vocabulary
  Rich, expressive text

━━━━━━━━━━━━━━━━━━━━━━

♻ PENALTY (Repetition Penalty)
Prevents word repetition.

• 1.0: No penalty
  May repeat words naturally

• 1.1-1.2: Moderate penalty
  Balanced, avoids repetition

• 1.3-1.5: Strong penalty
  Forces variety (may hurt quality)

━━━━━━━━━━━━━━━━━━━━━━

💡 TIPS:

✓ Technical docs: Temp 0.1, Top-P 0.5
✓ Fiction: Temp 0.7, Top-P 0.95
✓ Dialogue: Temp 0.4, Top-P 0.9
✓ Poetry: Temp 0.8, Top-P 1.0

⚠ High temp + low top-p = chaos
⚠ Low temp + high penalty = robotic"""
        else:
            help_content = """🌡 ТЕМПЕРАТУРА (Temperature)
Управляет креативностью и случайностью.

• 0.1-0.3: Очень точно, детерминировано
  Идеально для технического перевода

• 0.4-0.6: Сбалансировано, естественно
  Хорошо для обычного текста

• 0.7-1.2: Креативно, разнообразно
  Лучше для художественного перевода

━━━━━━━━━━━━━━━━━━━━━━

🎯 TOP-P (Nucleus Sampling)
Ограничивает разнообразие словаря.

• 0.1-0.5: Консервативные слова
  Строгий, формальный стиль

• 0.6-0.8: Сбалансированный словарь
  Естественный, повседневный язык

• 0.9-1.0: Полный словарь
  Богатый, выразительный текст

━━━━━━━━━━━━━━━━━━━━━━

♻ ШТРАФ (Repetition Penalty)
Предотвращает повторение слов.

• 1.0: Без штрафа
  Может повторять слова естественно

• 1.1-1.2: Умеренный штраф
  Сбалансировано, избегает повторов

• 1.3-1.5: Сильный штраф
  Заставляет разнообразие (может снизить качество)

━━━━━━━━━━━━━━━━━━━━━━

💡 СОВЕТЫ:

✓ Техническая документация: Temp 0.1, Top-P 0.5
✓ Художественная литература: Temp 0.7, Top-P 0.95
✓ Диалоги: Temp 0.4, Top-P 0.9
✓ Поэзия: Temp 0.8, Top-P 1.0

⚠ Высокая temp + низкий top-p = хаос
⚠ Низкая temp + высокий penalty = роботизированность"""

        help_text.insert("1.0", help_content)
        help_text.configure(state="disabled")

        self.refresh()

    def refresh(self):
        self.listbox.delete(0, "end")
        [self.listbox.insert("end", name) for name in self.prompts]

    def on_select(self, e):
        if not self.listbox.curselection(): return
        name = self.listbox.get(self.listbox.curselection())
        p_data = self.prompts[name]
        self.name_entry.delete(0, "end"); self.name_entry.insert(0, name)
        self.text_editor.delete("1.0", "end"); self.text_editor.insert("end", p_data["text"])
        self.s_temp.set(p_data["temp"]); self.l_temp.configure(text=f'{p_data["temp"]:.2f}')
        self.s_top.set(p_data["top_p"]); self.l_top.configure(text=f'{p_data["top_p"]:.2f}')
        self.s_pen.set(p_data["penalty"]); self.l_pen.configure(text=f'{p_data["penalty"]:.2f}')

    def add_p(self):
        self.name_entry.delete(0, "end"); self.name_entry.insert(0, "New Prompt")
        self.text_editor.delete("1.0", "end")
        self.s_temp.set(0.2); self.l_temp.configure(text="0.20")
        self.s_top.set(0.9); self.l_top.configure(text="0.90")
        self.s_pen.set(1.1); self.l_pen.configure(text="1.10")

    def del_p(self):
        name = self.name_entry.get()
        if name in self.prompts:
            del self.prompts[name]
            self.save_callback()
            self.refresh()

    def save_p(self):
        n, t = self.name_entry.get().strip(), self.text_editor.get("1.0", "end").strip()
        if n and t:
            self.prompts[n] = {
                "text": t,
                "temp": round(float(self.s_temp.get()), 2),
                "top_p": round(float(self.s_top.get()), 2),
                "penalty": round(float(self.s_pen.get()), 2)
            }
            self.save_callback()
            self.refresh()

# --- Главное приложение ---
# ============================================================================
# Main Application Class / Главный класс приложения
# Handles UI, translation logic, and AI engine management
# Управляет интерфейсом, логикой перевода и AI движком
# ============================================================================
class TranslatorApp(ctk.CTk):
    def __init__(self, model_path, gpu_layers, engine):
        super().__init__()
        self.model_path, self.gpu_layers, self.engine = model_path, gpu_layers, engine
        self.client, self.engine_proc = None, None
        self.is_translating = False

        # Translation cache for repeated phrases (performance optimization)
        # Кэш переводов для повторяющихся фраз (оптимизация производительности)
        self.translation_cache = OrderedDict()
        self.cache_max_size = 500

        # Model download cancellation flags / Флаги отмены загрузки моделей
        self.download_cancel_flags = {}

        self.config_ai = self.load_json(CONFIG_FILE, DEFAULT_CONFIG)
        self.ui_lang = self.config_ai.get("lang", "RU")

        # Load batch settings / Загрузка настроек батчинга
        self.batch_tokens = self.config_ai.get("batch_tokens", 600)
        self.batch_lines = self.config_ai.get("batch_lines", 8)
        self.smart_mode = self.config_ai.get("smart_mode", True)  # Smart mode by default / Умный режим по умолчанию

        # Load model from config or use provided / Загрузка модели из конфига или использование переданной
        if "model" in self.config_ai and os.path.exists(os.path.join(BASE_PATH, self.config_ai["model"])):
            self.model_path = self.config_ai["model"]

        self.load_and_migrate_prompts()

        self.title("ShallowL Translate AI Standalone")
        self.geometry("1300x850")
        ctk.set_appearance_mode(self.config_ai.get("theme", "dark"))

        icon_path = os.path.join(BASE_PATH, "icon.ico")
        if os.path.exists(icon_path): self.iconbitmap(icon_path)

        self.protocol("WM_DELETE_WINDOW", self.on_closing)
        self.setup_ui()
        self.apply_ui_lang()

        self.bind_all("<Control-KeyPress>", self._smart_hotkey_handler)
        self.lock_ui(True)

        # UI update queue and AI task queue / Очередь обновлений UI и задач AI
        self.ui_queue = queue.Queue(); self.process_queue()
        self.ai_task_queue = queue.Queue()
        threading.Thread(target=self.ai_thread_loop, daemon=True).start()

    def t(self, key, category="main_window"):
        """Get translation by key based on current UI language
        Получить перевод по ключу в зависимости от текущего языка

        Args:
            key: translation key / ключ перевода
            category: category (main_window, model_manager, batch_settings, etc.)
        """
        try:
            return TRANSLATIONS[self.ui_lang][category][key]
        except:
            # Fallback to English / Резерв на английский
            try:
                return TRANSLATIONS["EN"][category][key]
            except:
                return key

    def load_and_migrate_prompts(self):
        """Load prompts and migrate old format to new
        Загрузить промпты и мигрировать старый формат в новый
        """
        raw_prompts = self.load_json(PROMPTS_FILE, DEFAULT_PROMPTS)
        self.prompts = {}
        migrated = False
        for k, v in raw_prompts.items():
            if isinstance(v, str):
                if k in DEFAULT_PROMPTS: self.prompts[k] = DEFAULT_PROMPTS[k]
                else: self.prompts[k] = {"text": v, "temp": 0.2, "top_p": 0.9, "penalty": 1.1}
                migrated = True
            else: self.prompts[k] = v
        if migrated: self.save_json(PROMPTS_FILE, self.prompts)

    def _smart_hotkey_handler(self, event):
        """Handle Ctrl+C/V/X/A hotkeys for text widgets
        Обработка горячих клавиш Ctrl+C/V/X/A для текстовых виджетов
        """
        if event.keysym.lower() in ['c', 'v', 'x', 'a']: return
        kc = event.keycode; widget = event.widget
        try:
            if kc == 67: widget.event_generate("<<Copy>>"); return "break"
            elif kc == 86: widget.event_generate("<<Paste>>"); return "break"
            elif kc == 88: widget.event_generate("<<Cut>>"); return "break"
            elif kc == 65: 
                if hasattr(widget, "tag_add"): widget.tag_add("sel", "1.0", "end")
                elif hasattr(widget, "select_range"): widget.select_range(0, "end"); widget.icursor("end")
                return "break"
        except Exception: pass

    def setup_ui(self):
        self.grid_columnconfigure(0, weight=1); self.grid_columnconfigure(1, weight=1)
        self.grid_rowconfigure(1, weight=1)

        top_bar = ctk.CTkFrame(self, fg_color="transparent")
        top_bar.grid(row=0, column=0, columnspan=2, sticky="ew", padx=15, pady=10)
        top_bar.grid_columnconfigure(0, weight=1); top_bar.grid_columnconfigure(1, weight=1)

        ltop = ctk.CTkFrame(top_bar, fg_color="transparent")
        ltop.grid(row=0, column=0, sticky="ew", padx=(0,5))
        self.btn_load = ctk.CTkButton(ltop, text="Load", width=160, command=self.load_file)
        self.btn_load.pack(side="left")
        self.combo_src = ctk.CTkComboBox(ltop, values=SUPPORTED_LANGS, width=150)
        self.combo_src.set("English"); self.combo_src.pack(side="right")

        rtop = ctk.CTkFrame(top_bar, fg_color="transparent")
        rtop.grid(row=0, column=1, sticky="ew", padx=(5,0))
        self.combo_tgt = ctk.CTkComboBox(rtop, values=SUPPORTED_LANGS, width=150)
        self.combo_tgt.set("Russian"); self.combo_tgt.pack(side="left")
        self.btn_save = ctk.CTkButton(rtop, text="Save", width=160, command=self.save_file)
        self.btn_save.pack(side="right")

        self.text_l = ctk.CTkTextbox(self, font=("Arial", 14), wrap="word", border_width=1)
        self.text_l.grid(row=1, column=0, padx=(15, 5), pady=5, sticky="nsew")
        self.text_r = ctk.CTkTextbox(self, font=("Arial", 14), wrap="word", border_width=1)
        self.text_r.grid(row=1, column=1, padx=(5, 15), pady=5, sticky="nsew")

        ctrl_bar = ctk.CTkFrame(self, fg_color="transparent")
        ctrl_bar.grid(row=2, column=0, columnspan=2, sticky="ew", padx=15, pady=5)
        self.btn_go = ctk.CTkButton(ctrl_bar, text="▶ TRANSLATE", fg_color="#28a745", font=("Arial", 14, "bold"), height=45, command=self.start_trans)
        self.btn_go.pack(side="left", fill="x", expand=True, padx=(0,5))
        self.btn_stop = ctk.CTkButton(ctrl_bar, text="⏹ STOP", fg_color="#dc3545", font=("Arial", 14, "bold"), height=45, state="disabled", command=self.stop_trans)
        self.btn_stop.pack(side="right", fill="x", expand=True, padx=(5,0))

        bottom_row = ctk.CTkFrame(self, fg_color="transparent")
        bottom_row.grid(row=3, column=0, columnspan=2, sticky="ew", padx=15, pady=(5, 15))
        
        # Левый блок настроек интерфейса
        left_ui_settings = ctk.CTkFrame(bottom_row, fg_color="transparent")
        left_ui_settings.pack(side="left", anchor="s", padx=(0, 15))
        
        ctk.CTkLabel(left_ui_settings, text="UI Lang:", font=("Arial", 10)).pack(anchor="w")
        self.ui_lang_btn = ctk.CTkSegmentedButton(left_ui_settings, values=["EN", "RU"], width=80, command=self.change_ui_lang)
        self.ui_lang_btn.set(self.ui_lang); self.ui_lang_btn.pack(pady=(0, 8))

        ctk.CTkLabel(left_ui_settings, text="Theme:", font=("Arial", 10)).pack(anchor="w")
        self.theme_var = tk.StringVar(value=self.config_ai["theme"].capitalize())
        ctk.CTkRadioButton(left_ui_settings, text="Dark", variable=self.theme_var, value="Dark", command=self.change_theme).pack(pady=4, anchor="w")
        ctk.CTkRadioButton(left_ui_settings, text="Light", variable=self.theme_var, value="Light", command=self.change_theme).pack(pady=4, anchor="w")

        # Динамический цвет консоли
        log_frame = ctk.CTkFrame(bottom_row)
        log_frame.pack(side="left", fill="both", expand=True)

        self.log_a = ctk.CTkTextbox(log_frame, height=120, fg_color=("#e8e8e8", "#1e1e1e"),
                                     text_color=("#000000", "#00ff00"))
        self.log_a.pack(side="top", fill="both", expand=True, padx=5, pady=5)

        # Правый блок управления ИИ
        side_panel = ctk.CTkFrame(bottom_row, width=300)
        side_panel.pack(side="right", fill="y", padx=(15, 0))

        # Переключатель режима перевода
        mode_frame = ctk.CTkFrame(side_panel)
        mode_frame.pack(pady=(15, 10), padx=10, fill="x")

        self.lbl_translation_mode = ctk.CTkLabel(mode_frame, text=self.t('translation_mode'), font=("Arial", 10, "bold"))
        self.lbl_translation_mode.pack(pady=(5, 2))
        self.mode_var = tk.StringVar(value=self.t('mode_smart') if self.smart_mode else self.t('mode_fast'))
        self.mode_switch = ctk.CTkSegmentedButton(mode_frame, values=[self.t('mode_fast'), self.t('mode_smart')],
                                             variable=self.mode_var, command=self.change_mode)
        self.mode_switch.pack(pady=5)

        self.lbl_style = ctk.CTkLabel(side_panel, text="Style / Стиль:", font=("Arial", 12, "bold"))
        self.lbl_style.pack(pady=(10,0))
        self.combo_p = ctk.CTkComboBox(side_panel, values=list(self.prompts.keys()), width=240)
        self.combo_p.set(list(self.prompts.keys())[0]); self.combo_p.pack(pady=5)
        
        b_row = ctk.CTkFrame(side_panel, fg_color="transparent"); b_row.pack(pady=5)
        self.btn_pr = ctk.CTkButton(b_row, text=f"📝 {self.t('prompts', 'main_window')}", width=230, command=self.open_prompts)
        self.btn_pr.pack(padx=3)

        self.btn_models = ctk.CTkButton(side_panel, text=f"🤖 {self.t('title', 'model_manager')}", width=230,
                                        fg_color="#6c757d", command=self.open_model_manager)
        self.btn_models.pack(pady=5)

        self.btn_batch = ctk.CTkButton(side_panel, text=f"⚡ {self.t('title', 'batch_settings')}", width=230,
                                       fg_color="#17a2b8", command=self.open_batch_settings)
        self.btn_batch.pack(pady=5)

    def apply_ui_lang(self):
        """Применить локализацию ко всем элементам интерфейса"""
        # Главное окно
        self.title("ShallowL Translate AI Standalone")

        # Кнопки верхней панели
        self.btn_load.configure(text=self.t("load"))
        self.btn_save.configure(text=self.t("save"))

        # Кнопки управления переводом
        self.btn_go.configure(text=f"▶ {self.t('translate').upper()}")
        self.btn_stop.configure(text=f"⏹ {self.t('stop').upper()}")

        # Метки настроек
        if hasattr(self, 'lbl_theme'):
            self.lbl_theme.configure(text=self.t("theme") + ":")
        if hasattr(self, 'lbl_mode'):
            self.lbl_mode.configure(text=self.t("translation_mode") + ":")
        if hasattr(self, 'lbl_style'):
            self.lbl_style.configure(text=self.t("style") + ":")

        # Кнопки боковой панели
        if hasattr(self, 'btn_pr'):
            self.btn_pr.configure(text=f"📝 {self.t('prompts', 'main_window')}")
        if hasattr(self, 'btn_batch'):
            self.btn_batch.configure(text=f"⚡ {self.t('title', 'batch_settings')}")
        if hasattr(self, 'btn_models'):
            self.btn_models.configure(text=f"🤖 {self.t('title', 'model_manager')}")

        # Обновляем метку режима перевода
        if hasattr(self, 'lbl_translation_mode'):
            self.lbl_translation_mode.configure(text=self.t('translation_mode'))

        # Обновляем сегментированные кнопки режима
        if hasattr(self, 'mode_switch'):
            current_mode = self.mode_var.get()
            # Обновляем значения кнопок
            self.mode_switch.configure(values=[self.t('mode_fast'), self.t('mode_smart')])
            # Восстанавливаем выбранный режим
            if current_mode == "Fast" or current_mode == self.t('mode_fast'):
                self.mode_var.set(self.t('mode_fast'))
            else:
                self.mode_var.set(self.t('mode_smart'))

    def change_ui_lang(self, lang):
        self.ui_lang = lang
        self.config_ai["lang"] = lang
        self.save_json(CONFIG_FILE, self.config_ai)
        self.apply_ui_lang()

    def change_theme(self):
        t = self.theme_var.get().lower(); self.config_ai["theme"] = t; ctk.set_appearance_mode(t)
        self.save_json(CONFIG_FILE, self.config_ai)

    def lock_ui(self, lock):
        s = "disabled" if lock else "normal"
        # Кнопка моделей всегда доступна (чтобы можно было скачать модель)
        for w in [self.btn_load, self.btn_save, self.btn_go, self.combo_src, self.combo_tgt, self.combo_p, self.btn_pr, self.btn_batch]:
            if w: w.configure(state=s)

    def load_file(self):
        p = filedialog.askopenfilename(filetypes=[("Documents", "*.docx *.pdf *.txt")])
        if not p: return
        self.text_l.delete("1.0", "end"); ext = os.path.splitext(p)[1].lower()
        try:
            if ext == ".docx":
                doc = docx.Document(p); self.text_l.insert("end", "\n".join([pa.text for pa in doc.paragraphs]))
            elif ext == ".pdf":
                with fitz.open(p) as d: self.text_l.insert("end", "\n".join([pg.get_text() for pg in d]))
            else:
                with open(p, 'r', encoding='utf-8') as f: self.text_l.insert("end", f.read())
        except Exception as e: messagebox.showerror("Error", str(e))

    def save_file(self):
        p = filedialog.asksaveasfilename(defaultextension=".txt", filetypes=[("Text", "*.txt"), ("Word", "*.docx")])
        if not p: return
        txt = self.text_r.get("1.0", "end").strip(); ext = os.path.splitext(p)[1].lower()
        try:
            if ext == ".docx":
                d = docx.Document(); d.add_paragraph(txt); d.save(p)
            else:
                with open(p, 'w', encoding='utf-8') as f: f.write(txt)
        except Exception as e: messagebox.showerror("Error", str(e))

    # ========================================================================
    # AI Engine Management / Управление AI движком
    # ========================================================================

    def start_engine(self):
        """Start KoboldCPP AI engine with model
        Запуск AI движка KoboldCPP с моделью
        """
        self.log("⏳ Loading AI engine... / Загрузка ИИ движка...")
        exe = os.path.join(BASE_PATH, "bin", "koboldcpp.exe")
        model = os.path.join(BASE_PATH, self.model_path)

        # Check if model exists / Проверка существования модели
        if not os.path.exists(model):
            self.log(f"❌ Model not found! / Модель не найдена!\n📂 Path: {model}\n\n💡 Click '🤖 Models' button to download / Нажмите '🤖 Модели' для загрузки")
            return False

        # KoboldCPP engine configuration (optimized for 16GB VRAM)
        # Конфигурация движка KoboldCPP (оптимизировано для 16GB VRAM)
        cmd = [
            exe,
            "--model", model,
            "--gpulayers", str(self.gpu_layers),
            "--port", "5001",
            "--contextsize", "8192",
            "--blasbatchsize", "512"
        ]
        cmd.append("--usecublas" if self.engine == "cuda" else "--usevulkan")

        try:
            self.log(f"🚀 Starting KoboldCPP... / Запуск KoboldCPP...\n📝 Command: {os.path.basename(model)}")
            self.engine_proc = subprocess.Popen(cmd, creationflags=subprocess.CREATE_NO_WINDOW if os.name == 'nt' else 0)
        except Exception as e:
            self.log(f"❌ Failed to start engine! / Ошибка запуска движка!\n{e}")
            return False

        # Wait for engine to load (max 2 minutes) / Ожидание загрузки (макс 2 минуты)
        self.log("⏳ Waiting for AI to load into memory... / Ожидание загрузки в память...")
        for i in range(120):
            try:
                if requests.get("http://127.0.0.1:5001/api/v1/model", timeout=1).status_code == 200:
                    self.client = OpenAI(base_url="http://127.0.0.1:5001/v1", api_key="sk-123")
                    self.log("✅ AI Engine Ready! / Нейросеть готова к работе!\n🎯 You can start translating / Можно начинать перевод")
                    self.lock_ui(False)
                    # Hide console window on Windows / Скрыть консольное окно на Windows
                    try:
                        h = ctypes.WinDLL('kernel32').GetConsoleWindow()
                        if h != 0: ctypes.WinDLL('user32').ShowWindow(h, 0)
                    except: pass
                    return True
            except:
                time.sleep(1)

        self.log("❌ Timeout! AI engine failed to start / Таймаут! Движок не запустился")
        return False

    def ai_thread_loop(self):
        """Background thread for AI engine and translation tasks
        Фоновый поток для AI движка и задач перевода
        """
        if not self.start_engine():
            self.log("\n⚠ AI engine not started / ИИ движок не запущен\n💡 Check if model is downloaded / Проверьте, скачана ли модель")
            return
        while True:
            t = self.ai_task_queue.get()
            if t: self._do_translation(*t)

    # ========================================================================
    # Helper Methods / Вспомогательные методы
    # ========================================================================

    def _estimate_tokens(self, text):
        """Estimate token count (1 token ≈ 4 characters)
        Примерная оценка токенов (1 токен ≈ 4 символа)
        """
        return len(text) // 4

    def _get_cache_key(self, text, src, tgt, p_name):
        """Generate cache key for translation
        Генерация ключа для кэша перевода
        """
        return f"{src}:{tgt}:{p_name}:{text.strip()}"

    def _add_to_cache(self, key, value):
        """Add to cache with size limit (LRU)
        Добавление в кэш с ограничением размера (LRU)
        """
        if len(self.translation_cache) >= self.cache_max_size:
            self.translation_cache.popitem(last=False)  # Remove oldest / Удаляем самый старый
        self.translation_cache[key] = value

    def _batch_paragraphs(self, paragraphs):
        """Group paragraphs into batches for optimization
        Группировка параграфов в батчи для оптимизации
        """
        batches = []
        current_batch = []
        current_tokens = 0

        for p in paragraphs:
            p_tokens = self._estimate_tokens(p)

            # If paragraph too large - separate batch / Если параграф слишком большой - отдельный батч
            if p_tokens > self.batch_tokens:
                if current_batch:
                    batches.append(current_batch)
                    current_batch = []
                    current_tokens = 0
                batches.append([p])
                continue

            # Проверяем, влезет ли в текущий батч
            if (current_tokens + p_tokens > self.batch_tokens or
                len(current_batch) >= self.batch_lines):
                batches.append(current_batch)
                current_batch = [p]
                current_tokens = p_tokens
            else:
                current_batch.append(p)
                current_tokens += p_tokens

        if current_batch:
            batches.append(current_batch)

        return batches

    def change_mode(self, mode):
        """Переключение режима перевода"""
        # Проверяем режим по переведенному значению
        smart_mode_text = self.t('mode_smart')
        self.smart_mode = (mode == smart_mode_text or mode == "Smart")
        self.config_ai["smart_mode"] = self.smart_mode
        self.save_json(CONFIG_FILE, self.config_ai)

        if self.ui_lang == "EN":
            mode_text = "Smart mode: Analysis + Glossary + Context" if self.smart_mode else "Fast mode: Direct translation"
        else:
            mode_text = "Умный режим: Анализ + Глоссарий + Контекст" if self.smart_mode else "Быстрый режим: Прямой перевод"
        self.log(f"🔄 {mode_text}")

    # ========================================================================
    # Smart Translation Mode Methods / Методы умного режима перевода
    # Three-stage process: extract entities → create glossary → translate
    # Трёхэтапный процесс: извлечение сущностей → создание глоссария → перевод
    # ========================================================================

    def extract_entities(self, text, src, tgt):
        """Stage 1: Extract names, terms, abbreviations from text
        Этап 1: Извлечение имен, терминов, аббревиатур из текста
        """
        self.log("📋 Step 1/3: Extracting entities...")

        prompt = f"""Analyze this {src} text and extract:
1. Names (people, places, organizations)
2. Technical terms and specialized vocabulary
3. Abbreviations and acronyms
4. Repeated key phrases

Text:
{text[:2000]}

Output as JSON:
{{"names": [], "terms": [], "abbreviations": [], "key_phrases": []}}"""

        try:
            resp = self.client.chat.completions.create(
                model="local",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                max_tokens=500
            )

            result = resp.choices[0].message.content
            # Parse JSON response / Парсим JSON ответ
            import json
            entities = json.loads(result)
            return entities
        except Exception as e:
            self.log(f"⚠ Entity extraction failed: {e}")
            return {"names": [], "terms": [], "abbreviations": [], "key_phrases": []}

    def create_glossary(self, entities, src, tgt):
        """Stage 2: Create translation glossary for consistent terminology
        Этап 2: Создание глоссария переводов для консистентной терминологии
        """
        self.log("📚 Step 2/3: Creating glossary...")

        all_terms = []
        all_terms.extend(entities.get("names", []))
        all_terms.extend(entities.get("terms", []))
        all_terms.extend(entities.get("abbreviations", []))

        if not all_terms:
            return {}

        # Limit to 30 terms to avoid context overflow / Ограничиваем до 30 терминов
        all_terms = all_terms[:30]

        prompt = f"""Translate these terms from {src} to {tgt} consistently.
Keep names transliterated if appropriate.

Terms: {', '.join(all_terms)}

Output as JSON: {{"term1": "translation1", "term2": "translation2"}}"""

        try:
            resp = self.client.chat.completions.create(
                model="local",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.2,
                max_tokens=800
            )

            result = resp.choices[0].message.content
            import json
            glossary = json.loads(result)
            self.log(f"✅ Glossary created: {len(glossary)} terms")
            return glossary
        except Exception as e:
            self.log(f"⚠ Glossary creation failed: {e}")
            return {}

    def smart_translate_with_context(self, paragraphs, src, tgt, p_name, glossary):
        """Stage 3: Translate with context and glossary for better quality
        Этап 3: Перевод с контекстом и глоссарием для лучшего качества
        """
        self.log("🧠 Step 3/3: Smart translation with context...")

        p_data = self.prompts.get(p_name, DEFAULT_PROMPTS["Строгий / Strict"])
        sys_p = p_data["text"].replace("{source}", src).replace("{target}", tgt)
        t_temp, t_top, t_pen = p_data.get("temp", 0.2), p_data.get("top_p", 0.9), max(0.0, p_data.get("penalty", 1.1) - 1.0)

        # Format glossary string / Формируем строку глоссария
        glossary_str = "\n".join([f"- {k} → {v}" for k, v in list(glossary.items())[:20]])

        batches = self._batch_paragraphs(paragraphs)
        prev_context = ""

        for batch_idx, batch in enumerate(batches):
            if not self.is_translating:
                break

            processed_batch = []
            for p in batch:
                if not p.strip() or not TEXT_CHECK_REGEX.search(p):
                    self.ui_queue.put(("text", p + "\n" if p.strip() else "\n"))
                    continue
                processed_batch.append(p)

            if not processed_batch:
                continue

            # Формируем запрос с контекстом и глоссарием
            context_info = f"\nPrevious context: {prev_context}" if prev_context else ""
            glossary_info = f"\n\nGlossary (use these translations):\n{glossary_str}" if glossary_str else ""

            user_content = f"""Translate to {tgt}:{context_info}{glossary_info}

Text:
{chr(10).join(processed_batch)}"""

            try:
                resp = self.client.chat.completions.create(
                    model="local",
                    messages=[{"role": "system", "content": sys_p}, {"role": "user", "content": user_content}],
                    temperature=t_temp,
                    top_p=t_top,
                    presence_penalty=t_pen,
                    stream=True
                )

                current_translation = ""
                glossary_output = False  # Флаг для пропуска вывода глоссария

                for c in resp:
                    if not self.is_translating:
                        break
                    if c.choices[0].delta.content:
                        chunk = c.choices[0].delta.content
                        current_translation += chunk

                        # Пропускаем вывод глоссария в окно перевода
                        if not glossary_output:
                            # Ищем начало реального перевода (после глоссария)
                            if "DICTIONARY" in current_translation or "СЛОВАРЬ" in current_translation or "Glossary" in current_translation:
                                continue
                            # Если встретили перевод текста, начинаем выводить
                            if len(current_translation.strip()) > 0 and not any(x in current_translation for x in ["→", "term", "DICTIONARY", "СЛОВАРЬ", "Glossary"]):
                                glossary_output = True

                        if glossary_output:
                            self.ui_queue.put(("text", chunk))

                self.ui_queue.put(("text", "\n"))

                # Save last paragraph as context for next batch / Сохраняем последний абзац как контекст
                if current_translation:
                    prev_context = current_translation.strip()[-200:]  # Last 200 chars / Последние 200 символов

            except Exception as e:
                self.log(f"❌ Translation error: {e}")
                break

    # ========================================================================
    # Main Translation Methods / Основные методы перевода
    # ========================================================================

    def _do_translation(self, text, src, tgt, p_name):
        """Main translation function with mode selection
        Главная функция перевода с выбором режима
        """

        if self.smart_mode:
            # Smart mode: analysis → glossary → context-aware translation
            # Умный режим: анализ → глоссарий → перевод с контекстом
            self.log("🧠 Smart Translation Mode activated")

            paragraphs = text.split("\n")

            # Stage 1: Extract entities / Этап 1: Извлечение сущностей
            entities = self.extract_entities(text, src, tgt)

            # Stage 2: Create glossary / Этап 2: Создание глоссария
            glossary = self.create_glossary(entities, src, tgt)

            # Stage 3: Smart translation with context / Этап 3: Умный перевод с контекстом
            self.smart_translate_with_context(paragraphs, src, tgt, p_name, glossary)

        else:
            # Fast mode: direct batch translation / Быстрый режим: прямой перевод батчами
            self.log("⚡ Fast Translation Mode activated")
            self.fast_translate(text, src, tgt, p_name)

        self.is_translating = False
        self.ui_queue.put(("btn_state", "stopped"))
        self.log("✅ Translation completed! / Перевод завершен!")

    def fast_translate(self, text, src, tgt, p_name):
        """Fast translation mode (direct batch processing)
        Быстрый режим перевода (прямая обработка батчами)
        """
        p_data = self.prompts.get(p_name, DEFAULT_PROMPTS["Строгий / Strict"])
        sys_p = p_data["text"].replace("{source}", src).replace("{target}", tgt)
        t_temp, t_top, t_pen = p_data.get("temp", 0.2), p_data.get("top_p", 0.9), max(0.0, p_data.get("penalty", 1.1) - 1.0)

        paragraphs = text.split("\n")
        batches = self._batch_paragraphs(paragraphs)

        for batch in batches:
            if not self.is_translating:
                break

            processed_batch = []
            for p in batch:
                if not p.strip():
                    self.ui_queue.put(("text", "\n"))
                    continue

                if not TEXT_CHECK_REGEX.search(p):
                    self.ui_queue.put(("text", p + "\n"))
                    continue

                processed_batch.append(p)

            if not processed_batch:
                continue

            # Проверка кэша для одиночных параграфов
            if len(processed_batch) == 1:
                cache_key = self._get_cache_key(processed_batch[0], src, tgt, p_name)
                if cache_key in self.translation_cache:
                    self.ui_queue.put(("text", self.translation_cache[cache_key] + "\n"))
                    continue

            # Формирование запроса
            is_batch = len(processed_batch) > 1

            if is_batch:
                user_content = f"Translate each paragraph to {tgt}. Separate translations with newlines:\n\n" + "\n\n".join(processed_batch)
            else:
                user_content = f"Translate to {tgt}: {processed_batch[0]}"

            try:
                resp = self.client.chat.completions.create(
                    model="local",
                    messages=[{"role":"system","content":sys_p},{"role":"user","content":user_content}],
                    temperature=t_temp, top_p=t_top, presence_penalty=t_pen, stream=True
                )

                current_translation = ""
                buffer = ""

                for c in resp:
                    if not self.is_translating:
                        break
                    if c.choices[0].delta.content:
                        chunk = c.choices[0].delta.content
                        current_translation += chunk
                        buffer += chunk

                        # Фильтруем нумерацию на лету для батчей
                        if is_batch:
                            while True:
                                match = re.search(r'(^|\n)(\d+)\.\s+', buffer)
                                if match:
                                    before = buffer[:match.start()]
                                    if before:
                                        self.ui_queue.put(("text", before))
                                    if match.group(1):
                                        self.ui_queue.put(("text", "\n"))
                                    buffer = buffer[match.end():]
                                else:
                                    if len(buffer) > 5:
                                        output = buffer[:-5]
                                        self.ui_queue.put(("text", output))
                                        buffer = buffer[-5:]
                                    break
                        else:
                            self.ui_queue.put(("text", chunk))
                            buffer = ""

                # Выводим остатки буфера
                if is_batch and buffer:
                    cleaned = re.sub(r'^\d+\.\s+', '', buffer)
                    if cleaned:
                        self.ui_queue.put(("text", cleaned))

                # Сохранение в кэш для одиночных параграфов
                if not is_batch:
                    cache_key = self._get_cache_key(processed_batch[0], src, tgt, p_name)
                    self._add_to_cache(cache_key, current_translation.strip())

                self.ui_queue.put(("text", "\n"))

            except:
                break

        self.is_translating = False
        self.ui_queue.put(("btn_state", "stopped"))
        self.log("✅ Translation completed! / Перевод завершен!")

    def process_queue(self):
        batch_text = ""
        batch_count = 0
        max_batch = 10  # Накапливаем до 10 обновлений

        try:
            while True:
                m, d = self.ui_queue.get_nowait()
                if m == "text":
                    batch_text += d
                    batch_count += 1
                    if batch_count < max_batch:
                        continue  # Накапливаем еще
                    # Применяем накопленный текст
                    if batch_text:
                        self.text_r.insert("end", batch_text)
                        self.text_r.see("end")
                        batch_text = ""
                        batch_count = 0
                elif m == "log":
                    self.log_a.configure(state="normal")
                    current_text = self.log_a.get("1.0", "end").strip()
                    if current_text:
                        self.log_a.insert("end", "\n" + d)
                    else:
                        self.log_a.insert("end", d)
                    self.log_a.see("end")
                    # Не блокируем, чтобы можно было копировать
                elif m == "btn_state":
                    if d == "translating": self.btn_go.configure(state="disabled"); self.btn_stop.configure(state="normal")
                    else: self.btn_go.configure(state="normal"); self.btn_stop.configure(state="disabled")
                elif m == "clear": self.text_r.delete("1.0", "end")
        except queue.Empty:
            pass  # Очередь пуста - это нормально

        # Применяем остатки батча ВСЕГДА, даже если очередь опустела
        if batch_text:
            self.text_r.insert("end", batch_text)
            self.text_r.see("end")

        self.after(50, self.process_queue)  # Вернул 50мс для отзывчивости

    def log(self, m):
        logging.info(m)  # Логируем в файл
        self.ui_queue.put(("log", m))  # И в UI

    def load_json(self, f, d):
        try:
            if not os.path.exists(f): 
                with open(f, 'w', encoding='utf-8') as file: json.dump(d, file, indent=4, ensure_ascii=False); return d
            with open(f, 'r', encoding='utf-8') as file: return json.load(file)
        except: return d
    def save_json(self, f, d):
        with open(f, 'w', encoding='utf-8') as file: json.dump(d, file, indent=4, ensure_ascii=False)
        
    def start_trans(self):
        txt = self.text_l.get("1.0", "end").strip()
        if txt and self.client:
            self.is_translating = True
            self.ui_queue.put(("btn_state", "translating")); self.ui_queue.put(("clear", ""))
            self.ai_task_queue.put((txt, self.combo_src.get(), self.combo_tgt.get(), self.combo_p.get()))
            
    def stop_trans(self): self.is_translating = False
    def open_prompts(self): PromptManager(self, self.prompts, lambda: self.save_json(PROMPTS_FILE, self.prompts))
    
    def open_model_manager(self):
        """Менеджер загрузки моделей"""
        mm = ctk.CTkToplevel(self)
        mm.title(f"🤖 {self.t('title', 'model_manager')}")
        mm.geometry("900x720")
        mm.attributes("-topmost", True)

        # Заголовок
        ctk.CTkLabel(mm, text=self.t('select_model', 'model_manager'),
                     font=("Arial", 14, "bold")).pack(pady=(10, 5))

        models_dir = os.path.join(BASE_PATH, "models")
        os.makedirs(models_dir, exist_ok=True)

        # Скроллируемый фрейм для моделей
        scroll_frame = ctk.CTkScrollableFrame(mm, width=850, height=500)
        scroll_frame.pack(padx=20, pady=(5, 5), fill="both", expand=True)

        # Словарь для хранения виджетов прогресса
        self.download_widgets = {}

        for model_name, model_info in AVAILABLE_MODELS.items():
            # Пропускаем разделители
            if not model_info["file"]:
                # Переводим заголовки разделов
                section_text = model_name
                if "8GB" in model_name:
                    section_text = self.t('vram_8gb', 'model_manager')
                elif "6GB" in model_name:
                    section_text = self.t('vram_6gb', 'model_manager')
                elif "4GB" in model_name:
                    section_text = self.t('vram_4gb', 'model_manager')

                ctk.CTkLabel(scroll_frame, text=f"━━━ {section_text} ━━━", font=("Arial", 13, "bold"),
                            text_color=("#1f538d", "#4a9eff")).pack(fill="x", padx=5, pady=(8, 3))
                continue

            model_path = os.path.join(models_dir, model_info["file"])
            is_downloaded = os.path.exists(model_path)

            # Фрейм для каждой модели с местом под прогресс-бар
            mf = ctk.CTkFrame(scroll_frame, height=75)
            mf.pack(fill="x", padx=5, pady=1)
            mf.pack_propagate(False)

            # Левая часть - инфо
            left_frame = ctk.CTkFrame(mf, fg_color="transparent")
            left_frame.pack(side="left", fill="both", expand=True, padx=6, pady=2)

            # Название и статус
            status = "✅" if is_downloaded else "⬇"
            name_label = ctk.CTkLabel(left_frame, text=f"{status} {model_name}",
                        font=("Arial", 12, "bold"), anchor="w")
            name_label.pack(anchor="w", pady=(0, 1))

            # Инфо с переводом
            info_text = f"{model_info['size']} | {self.t('speed', 'model_manager')} {model_info['speed']} | {self.t('quality', 'model_manager')} {model_info['quality']}"
            ctk.CTkLabel(left_frame, text=info_text, font=("Arial", 10), anchor="w", text_color="gray").pack(anchor="w")

            # Прогресс-бар (скрыт по умолчанию)
            progress_frame = ctk.CTkFrame(left_frame, fg_color="transparent", height=20)
            progress_frame.pack(fill="x", pady=(2, 0))
            progress_frame.pack_forget()

            progress_bar = ctk.CTkProgressBar(progress_frame, width=300, height=12)
            progress_bar.pack(side="left", padx=(0, 5))
            progress_bar.set(0)

            status_label = ctk.CTkLabel(progress_frame, text="", font=("Arial", 10), width=100)
            status_label.pack(side="left")

            # Сохраняем виджеты для доступа при загрузке
            self.download_widgets[model_name] = {
                "progress_frame": progress_frame,
                "progress_bar": progress_bar,
                "status_label": status_label,
                "name_label": name_label,
                "model_frame": mf
            }

            # Правая часть - кнопки
            btn_frame = ctk.CTkFrame(mf, fg_color="transparent", width=120)
            btn_frame.pack(side="right", padx=5, pady=2)
            btn_frame.pack_propagate(False)

            if is_downloaded:
                def make_use_cmd(file):
                    return lambda: self.set_model(file, mm)
                def make_delete_cmd(path, name):
                    return lambda: self.delete_model(path, name, mm)

                ctk.CTkButton(btn_frame, text=self.t('use', 'model_manager'), width=110, height=28, font=("Arial", 11),
                             command=make_use_cmd(model_info["file"])).pack(pady=1)
                ctk.CTkButton(btn_frame, text=self.t('delete', 'model_manager'), width=110, height=28, font=("Arial", 11), fg_color="#dc3545",
                             command=make_delete_cmd(model_path, model_name)).pack(pady=1)
            else:
                def make_download_cmd(url, file, name):
                    return lambda: self.download_model_inline(url, file, name)

                download_btn = ctk.CTkButton(btn_frame, text=self.t('download', 'model_manager'), width=110, height=28, font=("Arial", 11), fg_color="#28a745",
                             command=make_download_cmd(model_info["url"], model_info["file"], model_name))
                download_btn.pack(pady=1)
                self.download_widgets[model_name]["download_btn"] = download_btn

        # Нижний фрейм с фиксированной высотой
        bottom_frame = ctk.CTkFrame(mm, fg_color="transparent", height=90)
        bottom_frame.pack(fill="x", padx=20, pady=(5, 10))
        bottom_frame.pack_propagate(False)

        # Легенда для звёзд
        ctk.CTkLabel(bottom_frame, text=self.t('legend', 'model_manager'), font=("Arial", 10), text_color="gray").pack(pady=(8, 3))

        # Текущая модель
        current_text = f"{self.t('current_model', 'model_manager')} {os.path.basename(self.model_path)}"
        ctk.CTkLabel(bottom_frame, text=current_text, font=("Arial", 12, "bold")).pack(pady=(3, 8))

    def set_model(self, model_file, window):
        """Установить модель как активную с горячей заменой"""
        new_model_path = f"models/{model_file}"

        # Проверяем, отличается ли от текущей
        if new_model_path == self.model_path:
            messagebox.showinfo("Info", "This model is already active / Эта модель уже активна")
            if window:
                window.destroy()
            return

        # Сохраняем в конфиг
        self.config_ai["model"] = new_model_path
        self.save_json(CONFIG_FILE, self.config_ai)

        if window:
            window.destroy()

        # Автоматический перезапуск
        self.log("🔄 Restarting program... / Перезапуск программы...")

        # Перезапуск через subprocess с правильными флагами
        python = sys.executable
        script = os.path.abspath(__file__)

        # Закрываем движок перед перезапуском
        if self.engine_proc:
            try:
                self.engine_proc.terminate()
            except:
                pass
        if os.name == 'nt':
            os.system("taskkill /f /im koboldcpp.exe >nul 2>&1")

        # Запускаем новый процесс
        if os.name == 'nt':
            # Используем START для запуска в новом окне, bat-файл удаляет сам себя
            batch_content = f'@echo off\nstart "" "{python}" "{script}"\ntimeout /t 1 /nobreak >nul\ndel "%~f0"\n'
            batch_file = os.path.join(BASE_PATH, "restart.bat")
            with open(batch_file, 'w') as f:
                f.write(batch_content)
            subprocess.Popen([batch_file], shell=True)
        else:
            subprocess.Popen([python, script] + sys.argv[1:])

        # Завершаем текущий процесс
        self.after(500, lambda: os._exit(0))

    def delete_model(self, model_path, model_name, window):
        """Удалить модель"""
        if messagebox.askyesno("Confirm", f"Удалить модель {model_name}?"):
            try:
                os.remove(model_path)
                messagebox.showinfo("Success", "Модель удалена!")
                window.destroy()
                self.open_model_manager()
            except Exception as e:
                messagebox.showerror("Error", f"Ошибка удаления: {e}")

    def download_model_inline(self, url, filename, model_name):
        """Загрузить модель с встроенным прогресс-баром"""
        models_dir = os.path.join(BASE_PATH, "models")
        filepath = os.path.join(models_dir, filename)

        widgets = self.download_widgets.get(model_name)
        if not widgets:
            return

        # Флаг отмены загрузки - ВСЕГДА создаем новый
        self.download_cancel_flags[model_name] = {"cancelled": False}
        cancel_flag = self.download_cancel_flags[model_name]

        # Показываем прогресс-бар
        widgets["progress_frame"].pack(fill="x", pady=(2, 0))
        widgets["progress_bar"].set(0)
        widgets["status_label"].configure(text=self.t('starting', 'model_manager'))

        # Меняем кнопку на "Stop"
        if "download_btn" in widgets:
            download_btn = widgets["download_btn"]
            def make_stop_cmd(name):
                return lambda: self.cancel_download(name)

            download_btn.configure(
                text=self.t('stop', 'model_manager'),
                fg_color="#dc3545",
                command=make_stop_cmd(model_name)
            )

        def download_thread():
            try:
                from huggingface_hub import hf_hub_download
                from tqdm import tqdm
                import traceback

                # Создаем кастомный класс tqdm для перехвата прогресса
                class GUIProgressBar(tqdm):
                    def __init__(self, *args, **kwargs):
                        super().__init__(*args, **kwargs)
                        self.gui_widgets = widgets
                        self.model_name = model_name

                    def update(self, n=1):
                        super().update(n)
                        if self.total and self.total > 0:
                            progress = self.n / self.total
                            current_mb = self.n / (1024 * 1024)
                            total_gb = self.total / (1024 * 1024 * 1024)

                            # Обновляем GUI через after
                            self.gui_widgets["progress_bar"].master.after(
                                0, lambda p=progress: self.gui_widgets["progress_bar"].set(p)
                            )
                            self.gui_widgets["status_label"].master.after(
                                0, lambda s=current_mb, t=total_gb:
                                self.gui_widgets["status_label"].configure(text=f"{s:.0f}/{t:.1f} GB")
                            )

                # Парсим URL правильно
                self.log(f"📥 Downloading {model_name}...")

                # URL формат: https://huggingface.co/bartowski/MODEL-NAME-GGUF/resolve/main/file.gguf
                # Нужно извлечь: repo_id = bartowski/MODEL-NAME-GGUF, filename = file.gguf

                # Убираем https://huggingface.co/
                url_without_domain = url.replace("https://huggingface.co/", "")
                # Разбиваем по /resolve/main/
                parts = url_without_domain.split("/resolve/main/")

                if len(parts) != 2:
                    raise Exception(f"Invalid URL format: {url}")

                repo_id = parts[0]  # bartowski/MODEL-NAME-GGUF
                hf_filename = parts[1]  # file.gguf

                self.after(0, lambda: widgets["status_label"].configure(text=self.t('starting', 'model_manager')))
                self.after(0, lambda: widgets["progress_bar"].set(0))

                self.log(f"⬇ Starting download to: {models_dir}")
                self.log(f"⬇ Repo: {repo_id}")
                self.log(f"⬇ File: {hf_filename}")

                # Проверка отмены перед началом
                if cancel_flag["cancelled"]:
                    self.log(f"❌ Download cancelled before start: {model_name}")
                    self.after(0, lambda: self.reset_download_button(model_name))
                    return

                self.after(0, lambda: widgets["status_label"].configure(text=self.t('downloading', 'model_manager')))
                self.log(f"🔄 Calling hf_hub_download...")

                # Загрузка файла с кастомным прогресс-баром
                downloaded_path = hf_hub_download(
                    repo_id=repo_id,
                    filename=hf_filename,
                    local_dir=models_dir,
                    local_dir_use_symlinks=False,
                    tqdm_class=GUIProgressBar
                )

                self.log(f"✅ hf_hub_download completed!")
                self.log(f"✅ Downloaded to: {downloaded_path}")

                # Проверка отмены после загрузки
                if cancel_flag["cancelled"]:
                    self.log(f"❌ Download cancelled: {model_name}")
                    if os.path.exists(downloaded_path):
                        try:
                            os.remove(downloaded_path)
                            self.log(f"🗑 Removed partial file: {downloaded_path}")
                        except Exception as del_err:
                            self.log(f"⚠ Could not remove file: {del_err}")
                    self.after(0, lambda: self.reset_download_button(model_name))
                    return

                # Проверяем размер
                if os.path.exists(downloaded_path):
                    file_size = os.path.getsize(downloaded_path) / (1024**3)
                    self.log(f"📊 File size: {file_size:.2f} GB")

                    if file_size < 0.1:
                        raise Exception(f"File too small: {file_size:.2f} GB")

                    # Обновляем UI
                    self.after(0, lambda: widgets["progress_bar"].set(1.0))
                    self.after(0, lambda: widgets["status_label"].configure(text=f"✅ {file_size:.2f} GB"))
                    self.after(0, lambda: widgets["name_label"].configure(text=f"✅ {model_name}"))

                    # Скрываем прогресс и меняем кнопки
                    def update_buttons():
                        widgets["progress_frame"].pack_forget()
                        if "download_btn" in widgets:
                            widgets["download_btn"].pack_forget()

                        btn_frame = widgets["download_btn"].master if "download_btn" in widgets else None
                        if btn_frame:
                            def make_use_cmd(f):
                                return lambda: self.set_model(f, None)
                            def make_delete_cmd(p, n):
                                return lambda: self.delete_model(p, n, None)

                            ctk.CTkButton(btn_frame, text="Use", width=100, height=28, font=("Arial", 10),
                                         command=make_use_cmd(filename)).pack(pady=1)
                            ctk.CTkButton(btn_frame, text="Delete", width=100, height=28, font=("Arial", 10), fg_color="#dc3545",
                                         command=make_delete_cmd(downloaded_path, model_name)).pack(pady=1)

                    self.after(2000, update_buttons)
                else:
                    raise Exception("Downloaded file not found")

            except Exception as e:
                if not cancel_flag["cancelled"]:
                    error_details = traceback.format_exc()
                    self.log(f"❌ ERROR downloading {model_name}:\n{error_details}")

                    self.after(0, lambda: widgets["progress_bar"].set(0))
                    self.after(0, lambda: widgets["status_label"].configure(text="❌ Error"))
                    self.after(0, lambda: messagebox.showerror("Error", f"Download error: {model_name}\nCheck console"))

                self.after(0, lambda: self.reset_download_button(model_name))

                if os.path.exists(filepath):
                    try:
                        os.remove(filepath)
                    except:
                        pass

        threading.Thread(target=download_thread, daemon=True).start()

    def cancel_download(self, model_name):
        """Отменить загрузку модели"""
        if model_name in self.download_cancel_flags:
            self.download_cancel_flags[model_name]["cancelled"] = True
            self.log(f"🛑 Cancelling download: {model_name}")
            # Сразу возвращаем кнопку в состояние Download
            self.reset_download_button(model_name)

    def reset_download_button(self, model_name):
        """Вернуть кнопку Download после отмены/ошибки"""
        widgets = self.download_widgets.get(model_name)
        if widgets and "download_btn" in widgets:
            download_btn = widgets["download_btn"]

            # Получаем URL и filename из AVAILABLE_MODELS
            for m_name, m_info in AVAILABLE_MODELS.items():
                if m_name == model_name:
                    def make_download_cmd(url, file, name):
                        return lambda: self.download_model_inline(url, file, name)

                    download_btn.configure(
                        text=self.t('download', 'model_manager'),
                        fg_color="#28a745",
                        command=make_download_cmd(m_info["url"], m_info["file"], model_name)
                    )
                    break

            # Скрываем прогресс-бар
            widgets["progress_frame"].pack_forget()
            widgets["status_label"].configure(text="")

    def open_batch_settings(self):
        """Настройки батчинга с автопресетами для моделей"""
        bs = ctk.CTkToplevel(self)
        bs.title(f"⚡ {self.t('title', 'batch_settings')}")
        bs.geometry("700x550")
        bs.attributes("-topmost", True)

        ctk.CTkLabel(bs, text=self.t('description', 'batch_settings'),
                     font=("Arial", 13, "bold")).pack(pady=15)

        # Информация о текущей модели
        current_model = os.path.basename(self.model_path)
        model_preset = MODEL_BATCH_PRESETS.get(current_model, {"tokens": 600, "lines": 8, "vram": "Unknown"})

        info_frame = ctk.CTkFrame(bs)
        info_frame.pack(fill="x", padx=30, pady=10)

        ctk.CTkLabel(info_frame, text=f"{self.t('current_model', 'batch_settings')} {current_model}",
                     font=("Arial", 11, "bold")).pack(pady=5)
        ctk.CTkLabel(info_frame, text=f"{self.t('recommended', 'batch_settings')}: {model_preset['vram']} VRAM",
                     font=("Arial", 10), text_color="gray").pack()

        # Batch Tokens
        tokens_frame = ctk.CTkFrame(bs)
        tokens_frame.pack(fill="x", padx=30, pady=10)

        ctk.CTkLabel(tokens_frame, text=self.t('tokens_per_batch', 'batch_settings'),
                     font=("Arial", 11, "bold")).pack(anchor="w", pady=5)
        ctk.CTkLabel(tokens_frame, text=self.t('info', 'batch_settings'),
                     font=("Arial", 9), text_color="gray").pack(anchor="w")

        tokens_slider = ctk.CTkSlider(tokens_frame, from_=200, to=1500, number_of_steps=26,
                                      command=lambda v: tokens_label.configure(text=f"{int(v)} (~{int(v*0.75)} words)"))
        tokens_slider.set(self.batch_tokens)
        tokens_slider.pack(fill="x", pady=5)

        tokens_label = ctk.CTkLabel(tokens_frame, text=f"{self.batch_tokens} (~{int(self.batch_tokens*0.75)} words)",
                                    font=("Arial", 10, "bold"))
        tokens_label.pack(pady=5)

        # Batch Lines
        lines_frame = ctk.CTkFrame(bs)
        lines_frame.pack(fill="x", padx=30, pady=10)

        ctk.CTkLabel(lines_frame, text=self.t('lines_per_batch', 'batch_settings'),
                     font=("Arial", 11, "bold")).pack(anchor="w", pady=5)
        ctk.CTkLabel(lines_frame, text=self.t('info', 'batch_settings'),
                     font=("Arial", 9), text_color="gray").pack(anchor="w")

        lines_slider = ctk.CTkSlider(lines_frame, from_=1, to=20, number_of_steps=19,
                                     command=lambda v: lines_label.configure(text=f"{int(v)} lines"))
        lines_slider.set(self.batch_lines)
        lines_slider.pack(fill="x", pady=5)

        lines_label = ctk.CTkLabel(lines_frame, text=f"{self.batch_lines} lines",
                                   font=("Arial", 10, "bold"))
        lines_label.pack(pady=5)

        # Пресеты
        preset_frame = ctk.CTkFrame(bs)
        preset_frame.pack(fill="x", padx=30, pady=15)

        ctk.CTkLabel(preset_frame, text="Quick Presets:", font=("Arial", 11, "bold")).pack(pady=5)

        presets_row = ctk.CTkFrame(preset_frame, fg_color="transparent")
        presets_row.pack(pady=5)

        def apply_preset(tokens, lines):
            tokens_slider.set(tokens)
            lines_slider.set(lines)
            tokens_label.configure(text=f"{tokens} (~{int(tokens*0.75)} words)")
            lines_label.configure(text=f"{lines} lines")

        ctk.CTkButton(presets_row, text=f"{self.t('quality', 'batch_presets')}\n(400/5)", width=100,
                     command=lambda: apply_preset(400, 5)).pack(side="left", padx=5)
        ctk.CTkButton(presets_row, text=f"{self.t('balanced', 'batch_presets')}\n(600/8)", width=100, fg_color="#28a745",
                     command=lambda: apply_preset(600, 8)).pack(side="left", padx=5)
        ctk.CTkButton(presets_row, text=f"{self.t('speed', 'batch_presets')}\n(1000/15)", width=100,
                     command=lambda: apply_preset(1000, 15)).pack(side="left", padx=5)
        ctk.CTkButton(presets_row, text=f"{self.t('auto', 'batch_presets')}\n({model_preset['tokens']}/{model_preset['lines']})", width=100, fg_color="#17a2b8",
                     command=lambda: apply_preset(model_preset['tokens'], model_preset['lines'])).pack(side="left", padx=5)

        # Предупреждение
        warning_frame = ctk.CTkFrame(bs, fg_color=("#fff3cd", "#664d03"))
        warning_frame.pack(fill="x", padx=30, pady=10)
        ctk.CTkLabel(warning_frame, text="⚠ Too high values may cause hallucinations or context loss!",
                     font=("Arial", 9, "bold"), text_color=("#856404", "#ffc107")).pack(pady=8)

        # Кнопки
        btn_frame = ctk.CTkFrame(bs, fg_color="transparent")
        btn_frame.pack(pady=15)

        def save_settings():
            self.batch_tokens = int(tokens_slider.get())
            self.batch_lines = int(lines_slider.get())
            self.config_ai["batch_tokens"] = self.batch_tokens
            self.config_ai["batch_lines"] = self.batch_lines
            self.save_json(CONFIG_FILE, self.config_ai)
            messagebox.showinfo(self.t('success', 'messages'),
                              f"{self.t('success', 'messages')}!" if self.ui_lang == "EN" else "Настройки сохранены!")
            bs.destroy()

        ctk.CTkButton(btn_frame, text=f"💾 {self.t('apply', 'batch_settings')}", width=150,
                     fg_color="#28a745", command=save_settings).pack(side="left", padx=5)
        ctk.CTkButton(btn_frame, text=self.t('close', 'batch_settings'), width=150,
                     command=bs.destroy).pack(side="left", padx=5)

    def on_closing(self):
        if self.engine_proc:
            try: self.engine_proc.terminate()
            except: pass
        if os.name == 'nt': os.system("taskkill /f /im koboldcpp.exe >nul 2>&1")
        self.destroy(); os._exit(0)

if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--model", type=str, default="models/dolphin-2.9.3-mistral-nemo-12b.Q4_K_M.gguf")
    parser.add_argument("--gpu-layers", type=int, default=99); parser.add_argument("--engine", type=str, default="cuda")
    args = parser.parse_args(); app = TranslatorApp(args.model, args.gpu_layers, args.engine); app.mainloop()